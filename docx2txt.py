import re
from pathlib import Path
import tkinter as tk
from tkinter import filedialog, messagebox

from docx import Document


def sanitize_filename(name: str) -> str:
    # Windows 文件名不能包含这些字符
    return re.sub(r'[\\/:*?"<>|]+', "_", name).strip() or "output"


def docx_to_text(doc_path: Path) -> str:
    """
    把 docx 转成纯文本：
    - 段落：逐段输出
    - 表格：逐行输出（单元格用 Tab 分隔）
    """
    doc = Document(str(doc_path))
    lines = []

    # 段落
    for p in doc.paragraphs:
        t = (p.text or "").rstrip()
        if t:
            lines.append(t)

    # 表格
    if doc.tables:
        if lines:
            lines.append("")  # 段落和表格之间空一行
        for ti, table in enumerate(doc.tables, 1):
            lines.append(f"[Table {ti}]")
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    # cell.text 可能包含多段，压成一行
                    cell_text = re.sub(r"\s+", " ", (cell.text or "").strip())
                    cells.append(cell_text)
                lines.append("\t".join(cells))
            lines.append("")  # 表格之间空一行

    # 末尾确保有换行
    return "\n".join(lines).strip() + "\n"


def main():
    root = tk.Tk()
    root.withdraw()

    # 选 docx（可多选）
    file_paths = filedialog.askopenfilenames(
        title="选择要转换的 DOCX 文件（可多选）",
        filetypes=[("Word 文档", "*.docx")],
    )
    if not file_paths:
        return

    # 选输出目录
    out_dir = filedialog.askdirectory(title="选择输出 TXT 的目录")
    if not out_dir:
        return

    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    ok = 0
    failed = []

    for fp in file_paths:
        docx_path = Path(fp)
        try:
            text = docx_to_text(docx_path)
            out_name = sanitize_filename(docx_path.stem) + ".txt"
            out_path = out_dir / out_name
            out_path.write_text(text, encoding="utf-8")
            ok += 1
        except Exception as e:
            failed.append((str(docx_path), repr(e)))

    # 结果提示
    msg = f"转换完成：成功 {ok} 个"
    if failed:
        msg += f"\n失败 {len(failed)} 个：\n" + "\n".join([f"- {p}: {err}" for p, err in failed[:8]])
        if len(failed) > 8:
            msg += f"\n... 还有 {len(failed) - 8} 个失败未展示"

    messagebox.showinfo("DOCX → TXT 批量转换", msg)


if __name__ == "__main__":
    main()